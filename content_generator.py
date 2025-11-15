import os
import sys
import time
import json
import re
import requests
from datetime import datetime
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING

# =============================================================
# Configuration (API unchanged)
# =============================================================
API_KEY = ""
MODEL = "deepseek-chat"
BASE_URL = "https://api.deepseek.com/v1/chat/completions"

TEMPERATURE = 0.35
MAX_TOKENS = 6000
MAX_RETRIES = 5
BACKOFF = 2.0

OUTPUT_DIR = "output_docs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# =============================================================
# Logging helper
# =============================================================

def log(msg: str):
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] {msg}")

# =============================================================
# Sanitizer
# =============================================================

def sanitize_file_name(name: str) -> str:
    name = re.sub(r"[^\w\s-]", "", name).strip().replace(" ", "_")
    return name[:180]

# =============================================================
# DOCX STYLE ENGINE — consistent palette, readable weights
# =============================================================

def apply_styles(doc: Document):
    """
    Visual system:
    - Blue: StrandTitle, SubStrandHeading
    - Bold black: UnitHeading, SectionHeading, ExerciseHeading, QuizHeading
    - Subtle color accents for teacher callouts
    - Body text justified with tidy spacing
    """

    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def ensure_style(name: str, fn):
        if name not in doc.styles:
            s = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            fn(s)

    # Top-level (blue)
    ensure_style("StrandTitle", lambda s: (
        setattr(s, "base_style", doc.styles["Normal"]),
        setattr(s.font, "size", Pt(22)),
        setattr(s.font, "bold", True),
        setattr(s.font.color, "rgb", RGBColor(0, 32, 96)),
        setattr(s.paragraph_format, "alignment", WD_ALIGN_PARAGRAPH.CENTER),
        setattr(s.paragraph_format, "space_before", Pt(12)),
        setattr(s.paragraph_format, "space_after", Pt(12)),
        setattr(s.paragraph_format, "keep_with_next", True)
    ))

    ensure_style("SubStrandHeading", lambda s: (
        setattr(s, "base_style", doc.styles["Normal"]),
        setattr(s.font, "size", Pt(16)),
        setattr(s.font, "bold", True),
        setattr(s.font.color, "rgb", RGBColor(0, 51, 153)),
        setattr(s.paragraph_format, "space_before", Pt(18)),
        setattr(s.paragraph_format, "space_after", Pt(6)),
        setattr(s.paragraph_format, "keep_with_next", True)
    ))

    # Inner headings (bold black)
    ensure_style("UnitHeading", lambda s: (
        setattr(s, "base_style", doc.styles["Normal"]),
        setattr(s.font, "size", Pt(14)),
        setattr(s.font, "bold", True),
        setattr(s.font.color, "rgb", RGBColor(0, 0, 0)),
        setattr(s.paragraph_format, "space_before", Pt(10)),
        setattr(s.paragraph_format, "space_after", Pt(4)),
        setattr(s.paragraph_format, "keep_with_next", True)
    ))

    ensure_style("SectionHeading", lambda s: (
        setattr(s, "base_style", doc.styles["Normal"]),
        setattr(s.font, "size", Pt(13)),
        setattr(s.font, "bold", True),
        setattr(s.font.color, "rgb", RGBColor(0, 0, 0)),
        setattr(s.paragraph_format, "space_before", Pt(10)),
        setattr(s.paragraph_format, "space_after", Pt(4)),
        setattr(s.paragraph_format, "keep_with_next", True)
    ))

    ensure_style("ExerciseHeading", lambda s: (
        setattr(s, "base_style", doc.styles["Normal"]),
        setattr(s.font, "size", Pt(13)),
        setattr(s.font, "bold", True),
        setattr(s.font.color, "rgb", RGBColor(0, 0, 0)),
        setattr(s.paragraph_format, "space_before", Pt(10)),
        setattr(s.paragraph_format, "space_after", Pt(4)),
        setattr(s.paragraph_format, "keep_with_next", True)
    ))

    ensure_style("QuizHeading", lambda s: (
        setattr(s, "base_style", doc.styles["Normal"]),
        setattr(s.font, "size", Pt(13)),
        setattr(s.font, "bold", True),
        setattr(s.font.color, "rgb", RGBColor(0, 0, 0)),
        setattr(s.paragraph_format, "space_before", Pt(10)),
        setattr(s.paragraph_format, "space_after", Pt(4)),
        setattr(s.paragraph_format, "keep_with_next", True)
    ))

    # Body text
    ensure_style("BodyTextClean", lambda s: (
        setattr(s, "base_style", doc.styles["Normal"]),
        setattr(s.paragraph_format, "line_spacing_rule", WD_LINE_SPACING.SINGLE),
        setattr(s.paragraph_format, "space_after", Pt(6))
    ))

    # Image description (grey italic)
    ensure_style("ImageDescription", lambda s: (
        setattr(s, "base_style", doc.styles["Normal"]),
        setattr(s.font, "italic", True),
        setattr(s.font.color, "rgb", RGBColor(120, 120, 120)),
        setattr(s.paragraph_format, "left_indent", Inches(0.3)),
        setattr(s.paragraph_format, "space_after", Pt(6))
    ))

    # Placeholder (red bold)
    ensure_style("PlaceholderText", lambda s: (
        setattr(s, "base_style", doc.styles["Normal"]),
        setattr(s.font, "bold", True),
        setattr(s.font.color, "rgb", RGBColor(150, 0, 0)),
        setattr(s.paragraph_format, "space_after", Pt(6))
    ))

    # Teacher callouts
    ensure_style("TeacherNote", lambda s: (
        setattr(s, "base_style", doc.styles["Normal"]),
        setattr(s.font, "bold", True),
        setattr(s.font.color, "rgb", RGBColor(90, 60, 0)),
        setattr(s.paragraph_format, "space_before", Pt(8)),
        setattr(s.paragraph_format, "space_after", Pt(6))
    ))

    ensure_style("Differentiation", lambda s: (
        setattr(s, "base_style", doc.styles["Normal"]),
        setattr(s.font, "bold", True),
        setattr(s.font.color, "rgb", RGBColor(0, 110, 0)),
        setattr(s.paragraph_format, "space_before", Pt(8)),
        setattr(s.paragraph_format, "space_after", Pt(6))
    ))

    ensure_style("AssessmentCallout", lambda s: (
        setattr(s, "base_style", doc.styles["Normal"]),
        setattr(s.font, "bold", True),
        setattr(s.font.color, "rgb", RGBColor(0, 0, 120)),
        setattr(s.paragraph_format, "space_before", Pt(8)),
        setattr(s.paragraph_format, "space_after", Pt(6))
    ))

    # Calibrate built-in list styles (not bold)
    for builtin in ["List Bullet", "List Number"]:
        if builtin in doc.styles:
            s = doc.styles[builtin]
            s.font.name = "Calibri"
            s.font.size = Pt(12)
            s.font.bold = False

# =============================================================
# Paragraph helpers
# =============================================================

def add_strand_title(doc: Document, text: str):
    return doc.add_paragraph(text, style="StrandTitle")

def add_substrand_heading(doc: Document, number: int, text: str):
    return doc.add_paragraph(f"{number}. {text}", style="SubStrandHeading")

def add_unit_heading(doc: Document, text: str):
    p = doc.add_paragraph(text, style="UnitHeading")
    p.paragraph_format.keep_with_next = True
    return p

def add_section_heading(doc: Document, text: str):
    p = doc.add_paragraph(text, style="SectionHeading")
    p.paragraph_format.keep_with_next = True
    return p

def add_exercise_heading(doc: Document, text: str):
    p = doc.add_paragraph(text, style="ExerciseHeading")
    p.paragraph_format.keep_with_next = True
    return p

def add_quiz_heading(doc: Document, text: str):
    p = doc.add_paragraph(text, style="QuizHeading")
    p.paragraph_format.keep_with_next = True
    return p

def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text, style="BodyTextClean")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_image_description(doc: Document, text: str):
    doc.add_paragraph(text, style="ImageDescription")

def add_placeholder(doc: Document, text: str):
    doc.add_paragraph(text, style="PlaceholderText")

def add_teacher_note(doc: Document, text: str):
    doc.add_paragraph(text, style="TeacherNote")

def add_differentiation(doc: Document, text: str):
    doc.add_paragraph(text, style="Differentiation")

def add_assessment_callout(doc: Document, text: str):
    doc.add_paragraph(text, style="AssessmentCallout")

def add_bullet_item(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.3 * level)
    return p

def add_manual_number_item(doc: Document, index: int, text: str, level: int = 0):
    """
    Manual numbering paragraph to guarantee restart per section.
    Creates hanging indent that aligns wrapped lines like Word lists.
    """
    p = doc.add_paragraph(style="BodyTextClean")
    base_indent = 0.3 * (level + 1)
    p.paragraph_format.left_indent = Inches(base_indent)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{index}. ")
    r.bold = False
    p.add_run(text)
    return p

# =============================================================
# Text cleaning
# =============================================================

def clean_model_output(text: str) -> str:
    # Remove markdown headers and code fences
    text = re.sub(r"^\s*#{1,6}.*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^```.*?$", "", text, flags=re.MULTILINE)
    # Strip bold/italic markers to prevent random emphasis
    text = text.replace("**", "").replace("__", "").replace("_", "")
    # Convert stray '*' bullets to '-'
    text = re.sub(r"^\s*\*\s+", "- ", text, flags=re.MULTILINE)
    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

# =============================================================
# Block parsing (heuristics)
# =============================================================

LIST_BULLET_PREFIX = re.compile(r"^\s*(?:[-•]|[\u2022])\s+")
LIST_NUMBER_PREFIX = re.compile(r"^\s*(\d+[\.)]|[a-zA-Z][\.)]|[ivxIVX]+\))\s+")

IMAGE_DESC_PATTERN = re.compile(r"\[IMAGE DESCRIPTION\]", re.IGNORECASE)
PLACEHOLDER_PATTERN = re.compile(r"^placeholder:", re.IGNORECASE)
TEACHER_NOTE_PATTERN = re.compile(r"^(teacher(?:'s)? note|note:)", re.IGNORECASE)
DIFFERENTIATION_PATTERN = re.compile(r"^differentiation", re.IGNORECASE)
ASSESSMENT_WORD_PATTERN = re.compile(r"^assessment\b", re.IGNORECASE)
UNIT_PATTERN = re.compile(r"^unit\s+\d+[:\.\-]?\s*", re.IGNORECASE)
EXERCISE_PATTERN = re.compile(r"^exercise\s+\d+\b", re.IGNORECASE)
QUIZ_PATTERN = re.compile(r"^quiz\b", re.IGNORECASE)
EXAMPLE_PATTERN = re.compile(r"^example\b", re.IGNORECASE)

SMALL_WORDS = {"of","the","to","and","in","on","for","a","an","at","with","from"}

def looks_like_title_case(line: str) -> bool:
    if len(line) > 80 or line.endswith("."):
        return False
    words = [w for w in re.split(r"\s+", line.strip()) if w]
    if not (1 <= len(words) <= 8):
        return False
    caps = 0
    for w in words:
        core = re.sub(r"[^\w]", "", w)
        if not core:
            continue
        if core.lower() in SMALL_WORDS:
            continue
        if core[0].isupper():
            caps += 1
    return caps >= 1

def detect_section_heading(line: str) -> Dict[str, Any]:
    """
    Returns {'text': str, 'kind': str} or {}.
    Kinds: unit, exercise, quiz, assessment, regular
    """
    s = line.strip()

    if UNIT_PATTERN.match(s):
        return {"text": s, "kind": "unit"}
    if EXERCISE_PATTERN.match(s):
        return {"text": s, "kind": "exercise"}
    if QUIZ_PATTERN.match(s):
        return {"text": s, "kind": "quiz"}
    if ASSESSMENT_WORD_PATTERN.match(s):
        return {"text": s, "kind": "assessment"}
    if EXAMPLE_PATTERN.match(s):
        return {"text": "Example", "kind": "regular"}
    if s.endswith(":"):
        return {"text": s[:-1].strip(), "kind": "regular"}
    if looks_like_title_case(s):
        return {"text": s, "kind": "regular"}
    return {}

def parse_blocks(raw: str, doc_type: str) -> List[Dict[str, Any]]:
    """
    Parse raw text into structured blocks:
      section_heading(kind), paragraph, bullet_list, numbered_list,
      image_description, placeholder, teacher_note
    Also: turn '- Where: National Museum' into a heading + paragraph.
    """
    lines = raw.splitlines()
    blocks: List[Dict[str, Any]] = []
    current_list: Dict[str, Any] = {}

    def flush_list():
        nonlocal current_list
        if current_list:
            blocks.append(current_list)
            current_list = {}

    for original_line in lines:
        line = original_line.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_list()
            continue

        # Special patterns
        if IMAGE_DESC_PATTERN.search(stripped):
            flush_list()
            blocks.append({"type": "image_description", "text": stripped})
            continue

        if PLACEHOLDER_PATTERN.search(stripped):
            flush_list()
            blocks.append({"type": "placeholder", "text": stripped})
            continue

        if doc_type.lower().startswith("teacher"):
            if TEACHER_NOTE_PATTERN.search(stripped):
                flush_list()
                blocks.append({"type": "teacher_note", "text": stripped})
                continue
            if DIFFERENTIATION_PATTERN.search(stripped):
                flush_list()
                blocks.append({"type": "section_heading", "text": "Differentiation", "kind": "regular"})
                continue
            if ASSESSMENT_WORD_PATTERN.search(stripped):
                flush_list()
                blocks.append({"type": "section_heading", "text": "Assessment", "kind": "assessment"})
                continue

        # Headings
        heading = detect_section_heading(stripped)
        if heading:
            flush_list()
            blocks.append({"type": "section_heading", **heading})
            continue

        # Lists
        bullet_match = LIST_BULLET_PREFIX.match(stripped)
        number_match = LIST_NUMBER_PREFIX.match(stripped)

        if bullet_match:
            content = LIST_BULLET_PREFIX.sub("", stripped).strip()
            # Promote label bullets 'Label: value' to heading + paragraph
            m = re.match(r"^([A-Z][A-Za-z ]{1,40}):\s*(.+)$", content)
            if m:
                flush_list()
                blocks.append({"type": "section_heading", "text": m.group(1), "kind": "regular"})
                blocks.append({"type": "paragraph", "text": m.group(2)})
                continue

            if current_list and current_list.get("type") == "bullet_list":
                current_list["items"].append({"text": content, "level": 0})
            else:
                flush_list()
                current_list = {"type": "bullet_list", "items": [{"text": content, "level": 0}]}
            continue

        if number_match:
            content = LIST_NUMBER_PREFIX.sub("", stripped).strip()
            if current_list and current_list.get("type") == "numbered_list":
                current_list["items"].append({"text": content, "level": 0})
            else:
                flush_list()
                current_list = {"type": "numbered_list", "items": [{"text": content, "level": 0}]}
            continue

        # Fallback paragraph
        flush_list()
        blocks.append({"type": "paragraph", "text": stripped})

    flush_list()
    return blocks

# =============================================================
# Grade tone
# =============================================================

def get_grade_tone(grade: str) -> str:
    g = int(grade)
    if g <= 6:
        return "clear, friendly language with concrete Kenyan and global everyday examples"
    elif g <= 9:
        return "explanatory and analytical writing with both Kenyan and global contexts and comparisons"
    else:
        return "precise, academic exposition with rigorous detail, including Kenyan and global applications"

# =============================================================
# Prompts — fully immersive (Kenya + global, relationships, depth)
# =============================================================

def student_prompt(grade, subject, strand, substrand, strand_outcomes):
    outcomes_text = ", ".join(strand_outcomes)
    tone = get_grade_tone(grade)
    return f"""
You are a professional CBC textbook writer for Grade {grade} {subject}.

Write fully immersive student textbook content for the sub-strand below. Do NOT include practice questions or exercises.
Do NOT include quiz or assessment content (that is handled separately). No markdown and no titles.

Quality requirements:
- Use {tone}.
- Break ideas, concepts, and terms down clearly and thoroughly.
- Make relationships explicit (cause-effect, compare-contrast, sequence, part-whole, prerequisite → advanced).
- Include both Kenyan and global context/examples.
- When appropriate for Grade {grade}, include formal notation, symbols, or precise terminology.

Structure (use plain text headings in this order):
Overview:
Key Ideas and Relationships:
Detailed Explanations:
Worked Example (Kenyan Context):
Worked Example (Global Context):
Applications and Real-Life Connections:
Vocabulary and Terms:
Common Misconceptions and Corrections:
Summary:
[IMAGE DESCRIPTION] lines wherever visuals would help.

Strand: {strand}
Sub-Strand: {substrand}
Strand Learning Outcomes: {outcomes_text}
"""

def teacher_prompt(grade, subject, strand, substrand, strand_outcomes):
    outcomes_text = ", ".join(strand_outcomes)
    tone = get_grade_tone(grade)
    return f"""
You are generating a TEACHER'S GUIDE for Grade {grade} {subject}.

Write deep, practical teacher guidance for the sub-strand below. Do NOT include student practice question sets.
No markdown and no titles.

Quality requirements:
- Use {tone}.
- Map relationships between key ideas (include a narrative “concept map”).
- Emphasize Kenyan and global contexts and cross-curricular links.
- Provide questioning strategies and checkpoints for understanding.
- Include differentiation (support and extension), materials, and assessment strategies with criteria (no item lists).

Structure (use plain text headings in this order):
Lesson Objectives (aligned to outcomes):
Prerequisites and Concept Relationships:
Lesson Flow (Engage, Explore, Explain, Elaborate, Evaluate):
Questioning and Checks for Understanding:
Differentiation:
Materials and Resources:
Assessment Strategies and Criteria:
Common Misconceptions and Remedies:
Cross-Curricular and Real-World Connections:
[IMAGE DESCRIPTION] suggestions where useful.

Strand: {strand}
Sub-Strand: {substrand}
Strand Learning Outcomes: {outcomes_text}
"""

# =============================================================
# API Request Handler (unchanged)
# =============================================================

def api_request(prompt: str) -> str:
    headers = {"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"}
    payload = {
        "model": MODEL,
        "messages": [{"role": "system", "content": prompt}],
        "temperature": TEMPERATURE,
        "max_tokens": MAX_TOKENS
    }

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            response = requests.post(BASE_URL, headers=headers, json=payload, timeout=600)
            if response.status_code == 200:
                data = response.json()
                return clean_model_output(data["choices"][0]["message"]["content"])
            if response.status_code in (429, 500, 502, 503, 504):
                log(f"API retry {attempt} due to {response.status_code}")
                time.sleep(BACKOFF * attempt)
            else:
                response.raise_for_status()
        except Exception as e:
            log(f"API error attempt {attempt}: {e}")
            time.sleep(BACKOFF * attempt)

    return ""

# =============================================================
# DOCX Builder — structured rendering with list reset & suppression
# =============================================================

def render_blocks(doc: Document, blocks: List[Dict[str, Any]], doc_type: str, current_substrand: str):
    """
    Writes structured blocks with appropriate styles.
    - Numbered lists are rendered manually to restart per section.
    - Quiz and Assessment sections are placeholders; content under them is skipped.
    - Exercise sections are suppressed entirely.
    """
    suppress_until_next_heading: Optional[str] = None  # 'quiz' | 'assessment' | 'exercise'

    for blk in blocks:
        t = blk["type"]

        if t == "section_heading":
            suppress_until_next_heading = None
            kind = blk.get("kind", "regular")
            text = blk.get("text", "")

            if kind == "unit":
                add_unit_heading(doc, text)
            elif kind == "exercise":
                add_exercise_heading(doc, text)  # optional visible heading
                add_placeholder(doc, "Placeholder: Practice questions are intentionally omitted (use Quiz/Assessment).")
                suppress_until_next_heading = "exercise"
            elif kind == "quiz":
                add_quiz_heading(doc, text)
                add_placeholder(doc, f"Placeholder: Quiz for sub-strand '{current_substrand}'.")
                suppress_until_next_heading = "quiz"
            elif kind == "assessment":
                add_section_heading(doc, "Assessment")
                add_placeholder(doc, f"Placeholder: Assessment for sub-strand '{current_substrand}'.")
                suppress_until_next_heading = "assessment"
            else:
                add_section_heading(doc, text)
            continue

        if suppress_until_next_heading in {"quiz", "assessment", "exercise"}:
            continue

        if t == "paragraph":
            add_body_paragraph(doc, blk["text"])
        elif t == "image_description":
            add_image_description(doc, blk["text"])
        elif t == "placeholder":
            add_placeholder(doc, blk["text"])
        elif t == "teacher_note":
            add_teacher_note(doc, blk["text"])
        elif t == "bullet_list":
            for item in blk["items"]:
                add_bullet_item(doc, item["text"], item.get("level", 0))
        elif t == "numbered_list":
            for i, item in enumerate(blk["items"], start=1):
                add_manual_number_item(doc, i, item["text"], item.get("level", 0))
        else:
            if "text" in blk and blk["text"]:
                add_body_paragraph(doc, blk["text"])

def build_doc(grade: str, subject: str, strand_name: str, content_map: Dict[str, str], doc_type: str = "Student") -> Document:
    doc = Document()
    apply_styles(doc)

    # Exact strand name in title
    add_strand_title(doc, f"Grade {grade} {subject} - {strand_name} ({doc_type} Textbook)")
    add_body_paragraph(doc, f"Generated on {datetime.now().strftime('%Y-%m-%d')}")

    items = list(content_map.items())
    total = len(items)

    for idx, (sub, text) in enumerate(items, start=1):
        add_substrand_heading(doc, idx, sub)

        blocks = parse_blocks(text, doc_type)
        render_blocks(doc, blocks, doc_type, current_substrand=sub)

        if idx < total:
            doc.add_page_break()

    return doc

# =============================================================
# Main
# =============================================================

def main():
    if not API_KEY:
        log("Missing API KEY")
        sys.exit(1)

    with open("content.json", "r", encoding="utf-8") as f:
        curriculum = json.load(f)

    grade = "4"
    subject = "English"

    strand_name = list(curriculum[grade][subject].keys())[0]          # exact strand name
    strand_subs = list(curriculum[grade][subject][strand_name].items())[:2]

    log(f"Testing: Grade {grade}, Subject {subject}, Strand {strand_name}")

    strand_outcomes: List[str] = []
    for _, details in curriculum[grade][subject][strand_name].items():
        strand_outcomes.extend(details.get("learning_outcomes", []))

    student_map: Dict[str, str] = {}
    teacher_map: Dict[str, str] = {}

    for sub, _details in strand_subs:
        log(f"Generating Student → {sub}")
        s_prompt = student_prompt(grade, subject, strand_name, sub, strand_outcomes)
        s_content = api_request(s_prompt)
        if "quiz" not in s_content.lower():
            s_content += f"\n\nPlaceholder: Quiz for sub-strand '{sub}'"
        student_map[sub] = s_content

        log(f"Generating Teacher → {sub}")
        t_prompt = teacher_prompt(grade, subject, strand_name, sub, strand_outcomes)
        t_content = api_request(t_prompt)
        teacher_map[sub] = t_content

    # Strand-level assessment placeholders
    student_map[f"{strand_name} Assessment"] = f"Placeholder: Assessment for strand '{strand_name}'."
    teacher_map[f"{strand_name} Assessment"] = f"Placeholder: Assessment guidance for strand '{strand_name}'."

    s_doc = build_doc(grade, subject, strand_name, student_map, "Student")
    t_doc = build_doc(grade, subject, strand_name, teacher_map, "Teacher")

    s_file = os.path.join(OUTPUT_DIR, f"Grade{grade}_{subject}_{sanitize_file_name(strand_name)}_Student.docx")
    t_file = os.path.join(OUTPUT_DIR, f"Grade{grade}_{subject}_{sanitize_file_name(strand_name)}_Teacher.docx")

    s_doc.save(s_file)
    t_doc.save(t_file)

    log(f"Saved Student: {s_file}")
    log(f"Saved Teacher: {t_file}")

if __name__ == "__main__":
    main()